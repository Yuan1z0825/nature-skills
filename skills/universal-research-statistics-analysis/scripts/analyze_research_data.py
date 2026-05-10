#!/usr/bin/env python3
"""Universal Research Statistics Analysis Skill.

Generates an auditable Excel workbook, Word three-line tables, Methods/Results
text, optional figures, and logs from a CSV/XLSX research dataset.

Default policy:
- p-values only; no q-values/FDR columns by default.
- mean ± SEM in Word tables for continuous outcomes.
- median (IQR) for score/ordinal outcomes.
- compact letter display for >=3 groups; detailed post hoc hidden unless requested.
"""
from __future__ import annotations

import argparse
import json
import math
import re
import sys
import warnings
from copy import copy
from dataclasses import dataclass
from pathlib import Path
from typing import Dict, Iterable, List, Optional, Sequence, Tuple

import numpy as np
import pandas as pd
from scipy import stats

try:
    import matplotlib.pyplot as plt
except Exception:  # pragma: no cover
    plt = None

try:
    import statsmodels.api as sm
    import statsmodels.formula.api as smf
    from statsmodels.stats.anova import anova_lm
except Exception:  # pragma: no cover
    sm = None
    smf = None
    anova_lm = None

try:
    from docx import Document
    from docx.enum.section import WD_SECTION_START
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Cm, Pt
except Exception:  # pragma: no cover
    Document = None

SUPERSCRIPTS = str.maketrans({
    "a": "ᵃ", "b": "ᵇ", "c": "ᶜ", "d": "ᵈ", "e": "ᵉ", "f": "ᶠ", "g": "ᵍ",
    "h": "ʰ", "i": "ⁱ", "j": "ʲ", "k": "ᵏ", "l": "ˡ", "m": "ᵐ", "n": "ⁿ",
    "o": "ᵒ", "p": "ᵖ", "r": "ʳ", "s": "ˢ", "t": "ᵗ", "u": "ᵘ", "v": "ᵛ",
    "w": "ʷ", "x": "ˣ", "y": "ʸ", "z": "ᶻ"
})

ORDINAL_PATTERNS = [
    r"score", r"scale", r"grade", r"rank", r"rating", r"likert", r"stage", r"class"
]
ID_PATTERNS = [r"(^|_)id($|_)", r"subject", r"participant", r"animal", r"sample", r"record"]
DESIGN_NAMES = {"group", "treatment", "arm", "time", "visit", "day", "week", "period", "sequence", "carryover"}


def canonical_name(name: str) -> str:
    s = str(name).strip()
    s = re.sub(r"[\n\r\t]+", " ", s)
    s = re.sub(r"\s+", " ", s)
    return s


def resolve_col(df: pd.DataFrame, requested: Optional[str]) -> Optional[str]:
    if requested is None:
        return None
    if requested in df.columns:
        return requested
    target = normalize_key(requested)
    mapping = {normalize_key(c): c for c in df.columns}
    return mapping.get(target)


def normalize_key(s: str) -> str:
    return re.sub(r"[^a-z0-9]+", "", str(s).lower())


def is_categorical_label_series(s: pd.Series) -> bool:
    vals = s.dropna().astype(str).head(50)
    return any(bool(re.search(r"[A-Za-z]", v)) for v in vals)


def coerce_numeric_value(x):
    if pd.isna(x):
        return np.nan
    if isinstance(x, (int, float, np.integer, np.floating)) and not isinstance(x, bool):
        return float(x)
    s = str(x).strip()
    if s == "":
        return np.nan
    # Keep categorical labels such as A30, P40, T1, Dose2 as non-numeric.
    if re.search(r"[A-Za-z]", s):
        return np.nan
    s = s.replace("≤", "").replace("≥", "").replace("<", "").replace(">", "")
    # comma decimal but not thousands grouping
    if re.match(r"^-?\d+,\d+$", s):
        s = s.replace(",", ".")
    s = s.replace("%", "")
    try:
        return float(s)
    except Exception:
        return np.nan


def load_data(path: Path) -> pd.DataFrame:
    if path.suffix.lower() == ".csv":
        try:
            df = pd.read_csv(path)
        except UnicodeDecodeError:
            df = pd.read_csv(path, encoding="utf-8-sig")
    elif path.suffix.lower() in {".xlsx", ".xls"}:
        df = pd.read_excel(path)
    else:
        raise ValueError(f"Unsupported input format: {path.suffix}")
    df.columns = [canonical_name(c) for c in df.columns]
    return df


def numeric_coercion(df: pd.DataFrame, protected_cols: Iterable[str]) -> Tuple[pd.DataFrame, pd.DataFrame]:
    out = df.copy()
    protected = {c for c in protected_cols if c in out.columns}
    logs = []
    for col in out.columns:
        if col in protected:
            logs.append({"column": col, "action": "protected", "numeric_count": None, "non_missing": int(out[col].notna().sum())})
            continue
        if pd.api.types.is_numeric_dtype(out[col]):
            logs.append({"column": col, "action": "already_numeric", "numeric_count": int(out[col].notna().sum()), "non_missing": int(out[col].notna().sum())})
            continue
        if is_categorical_label_series(out[col]):
            logs.append({"column": col, "action": "kept_categorical_label", "numeric_count": 0, "non_missing": int(out[col].notna().sum())})
            continue
        coerced = out[col].map(coerce_numeric_value)
        n_num = int(coerced.notna().sum())
        n_nonmiss = int(out[col].notna().sum())
        if n_num > 0 and n_num / max(n_nonmiss, 1) >= 0.7:
            out[col] = coerced
            logs.append({"column": col, "action": "coerced_to_numeric", "numeric_count": n_num, "non_missing": n_nonmiss})
        else:
            logs.append({"column": col, "action": "kept_non_numeric", "numeric_count": n_num, "non_missing": n_nonmiss})
    return out, pd.DataFrame(logs)


def is_ordinal_like(col: str) -> bool:
    k = normalize_key(col)
    return any(re.search(p, k) for p in ORDINAL_PATTERNS)


def is_id_like(col: str) -> bool:
    k = normalize_key(col)
    return any(re.search(p, k) for p in ID_PATTERNS)


def numeric_outcomes(df: pd.DataFrame, exclude: Iterable[str]) -> List[str]:
    excluded = {c for c in exclude if c}
    outcomes = []
    for col in df.columns:
        if col in excluded:
            continue
        if is_id_like(col):
            continue
        if normalize_key(col) in DESIGN_NAMES:
            continue
        if pd.api.types.is_numeric_dtype(df[col]) and df[col].notna().sum() >= 3:
            if df[col].nunique(dropna=True) > 1:
                outcomes.append(col)
    return outcomes


def format_p(p: float) -> str:
    if pd.isna(p):
        return "NA"
    if p < 0.001:
        return "≤0.001"
    return f"{p:.3f}"


def format_num(x: float, digits: int = 2) -> str:
    if pd.isna(x):
        return "NA"
    return f"{x:.{digits}f}"


def sem(x: pd.Series) -> float:
    x = x.dropna().astype(float)
    if len(x) <= 1:
        return np.nan
    return float(x.std(ddof=1) / math.sqrt(len(x)))


def normal_enough(groups: List[pd.Series]) -> bool:
    ok = True
    for g in groups:
        g = g.dropna().astype(float)
        if len(g) < 3:
            return False
        if len(g) >= 3:
            try:
                # Shapiro can be overly strict in large samples; cap with a practical guard.
                stat, p = stats.shapiro(g.sample(min(len(g), 500), random_state=1) if len(g) > 500 else g)
                ok = ok and (p > 0.05)
            except Exception:
                ok = False
    return ok


def descriptive_by_group(df: pd.DataFrame, outcomes: List[str], group_col: Optional[str]) -> pd.DataFrame:
    rows = []
    if group_col and group_col in df.columns:
        groups = [g for g in pd.unique(df[group_col].dropna())]
        for var in outcomes:
            for g in groups:
                x = df.loc[df[group_col] == g, var].dropna().astype(float)
                rows.append({
                    "Variable": var, "Group": g, "n": int(len(x)),
                    "Mean": x.mean() if len(x) else np.nan,
                    "SD": x.std(ddof=1) if len(x) > 1 else np.nan,
                    "SEM": sem(x),
                    "Median": x.median() if len(x) else np.nan,
                    "Q1": x.quantile(0.25) if len(x) else np.nan,
                    "Q3": x.quantile(0.75) if len(x) else np.nan,
                    "Min": x.min() if len(x) else np.nan,
                    "Max": x.max() if len(x) else np.nan,
                })
    else:
        for var in outcomes:
            x = df[var].dropna().astype(float)
            rows.append({
                "Variable": var, "Group": "All", "n": int(len(x)),
                "Mean": x.mean() if len(x) else np.nan,
                "SD": x.std(ddof=1) if len(x) > 1 else np.nan,
                "SEM": sem(x),
                "Median": x.median() if len(x) else np.nan,
                "Q1": x.quantile(0.25) if len(x) else np.nan,
                "Q3": x.quantile(0.75) if len(x) else np.nan,
                "Min": x.min() if len(x) else np.nan,
                "Max": x.max() if len(x) else np.nan,
            })
    return pd.DataFrame(rows)


def pairwise_p(groups: Dict[str, pd.Series], parametric: bool) -> pd.DataFrame:
    rows = []
    names = list(groups.keys())
    for i in range(len(names)):
        for j in range(i + 1, len(names)):
            a, b = names[i], names[j]
            x = groups[a].dropna().astype(float)
            y = groups[b].dropna().astype(float)
            p = np.nan
            test = None
            if len(x) >= 2 and len(y) >= 2 and x.nunique() > 1 and y.nunique() > 1:
                if parametric:
                    _, p = stats.ttest_ind(x, y, equal_var=False, nan_policy="omit")
                    test = "Welch pairwise t-test"
                else:
                    _, p = stats.mannwhitneyu(x, y, alternative="two-sided")
                    test = "Pairwise Mann–Whitney U"
            rows.append({"Group 1": a, "Group 2": b, "Test": test, "p-value": p, "Significant": bool(pd.notna(p) and p < 0.05)})
    return pd.DataFrame(rows)


def compact_letters(group_stats: pd.DataFrame, pairs: pd.DataFrame, overall_p: float) -> Dict[str, str]:
    groups = list(group_stats["Group"])
    if len(groups) < 3:
        return {g: "" for g in groups}
    if pd.isna(overall_p) or overall_p >= 0.05:
        return {g: "a" for g in groups}
    sig = {tuple(sorted([r["Group 1"], r["Group 2"]])): bool(r["Significant"]) for _, r in pairs.iterrows()}
    # Sort high-to-low by mean for intuitive a>b ordering.
    means = dict(zip(group_stats["Group"], group_stats["Mean"]))
    ordered = sorted(groups, key=lambda g: (pd.isna(means.get(g)), -means.get(g, -np.inf) if not pd.isna(means.get(g)) else 0))
    letters: List[Tuple[str, List[str]]] = []
    alpha = list("abcdefghijklmnopqrstuvwxyz")
    assigned = {g: "" for g in groups}
    for g in ordered:
        placed = False
        for letter, holders in letters:
            conflict = any(sig.get(tuple(sorted([g, h])), False) for h in holders)
            if not conflict:
                holders.append(g)
                assigned[g] += letter
                placed = True
        if not placed:
            letter = alpha[len(letters)] if len(letters) < len(alpha) else f"l{len(letters)+1}"
            letters.append((letter, [g]))
            assigned[g] += letter
    return assigned


def group_comparisons(df: pd.DataFrame, outcomes: List[str], group_col: Optional[str]) -> Tuple[pd.DataFrame, pd.DataFrame, pd.DataFrame]:
    comp_rows, posthoc_rows, table_rows = [], [], []
    if not group_col or group_col not in df.columns:
        return pd.DataFrame(), pd.DataFrame(), pd.DataFrame()
    group_levels = [g for g in pd.unique(df[group_col].dropna())]
    if len(group_levels) < 2:
        return pd.DataFrame(), pd.DataFrame(), pd.DataFrame()
    for var in outcomes:
        groups = {str(g): df.loc[df[group_col] == g, var].dropna().astype(float) for g in group_levels}
        valid = {g: x for g, x in groups.items() if len(x) >= 2 and x.nunique() > 1}
        if len(valid) < 2:
            comp_rows.append({"Variable": var, "Test": None, "p-value": np.nan, "Reason": "Fewer than two valid groups with variance"})
            continue
        ordinal = is_ordinal_like(var)
        parametric = (not ordinal) and normal_enough(list(valid.values())) and all(len(x) >= 3 for x in valid.values())
        p = np.nan
        stat = np.nan
        if len(valid) == 2:
            g1, g2 = list(valid.keys())
            if parametric:
                stat, p = stats.ttest_ind(valid[g1], valid[g2], equal_var=False, nan_policy="omit")
                test = "Welch t-test"
            else:
                stat, p = stats.mannwhitneyu(valid[g1], valid[g2], alternative="two-sided")
                test = "Mann–Whitney U"
        else:
            if parametric:
                stat, p = stats.f_oneway(*list(valid.values()))
                test = "One-way ANOVA"
            else:
                stat, p = stats.kruskal(*list(valid.values()))
                test = "Kruskal–Wallis"
        comp_rows.append({"Variable": var, "Test": test, "Statistic": stat, "p-value": p, "Reason": "OK", "Parametric route": parametric, "Ordinal/score": ordinal})
        desc = descriptive_by_group(df, [var], group_col)
        desc_var = desc[desc["Variable"] == var].copy()
        if len(valid) >= 3:
            pairs = pairwise_p(valid, parametric=parametric)
            pairs.insert(0, "Variable", var)
            posthoc_rows.extend(pairs.to_dict("records"))
            letters = compact_letters(desc_var, pairs, p)
        else:
            letters = {str(g): "" for g in group_levels}
        row = {"Variable": var, "p-value": p, "Method": test, "Summary type": "median (IQR)" if ordinal else "mean ± SEM"}
        for _, r in desc_var.iterrows():
            g = str(r["Group"])
            letter = letters.get(g, "")
            supers = letter.translate(SUPERSCRIPTS) if letter else ""
            if ordinal:
                cell = f"{format_num(r['Median'])} ({format_num(r['Q1'])}–{format_num(r['Q3'])}){supers}"
            else:
                cell = f"{format_num(r['Mean'])} ± {format_num(r['SEM'])}{supers}"
            row[g] = cell
            row[f"{g} letter"] = letter
        table_rows.append(row)
    return pd.DataFrame(comp_rows), pd.DataFrame(posthoc_rows), pd.DataFrame(table_rows)


def paired_comparisons(df: pd.DataFrame, outcomes: List[str], subject_col: Optional[str], time_col: Optional[str]) -> pd.DataFrame:
    if not subject_col or not time_col or subject_col not in df.columns or time_col not in df.columns:
        return pd.DataFrame()
    levels = list(pd.unique(df[time_col].dropna()))
    if len(levels) != 2:
        return pd.DataFrame([{"Reason": "Paired comparison requires exactly two time levels"}])
    rows = []
    t1, t2 = levels[0], levels[1]
    for var in outcomes:
        wide = df[[subject_col, time_col, var]].dropna().pivot_table(index=subject_col, columns=time_col, values=var, aggfunc="mean")
        if t1 not in wide.columns or t2 not in wide.columns:
            continue
        paired = wide[[t1, t2]].dropna()
        if len(paired) < 3:
            rows.append({"Variable": var, "Test": None, "p-value": np.nan, "Reason": "n<3 paired observations"})
            continue
        diff = paired[t2] - paired[t1]
        if (not is_ordinal_like(var)) and len(diff) >= 3 and stats.shapiro(diff).pvalue > 0.05:
            stat, p = stats.ttest_rel(paired[t2], paired[t1])
            test = "Paired t-test"
        else:
            stat, p = stats.wilcoxon(paired[t2], paired[t1])
            test = "Wilcoxon signed-rank"
        rows.append({"Variable": var, "Time 1": t1, "Time 2": t2, "n pairs": len(paired), "Test": test, "Statistic": stat, "p-value": p, "Mean change": diff.mean(), "Reason": "OK"})
    return pd.DataFrame(rows)


def safe_formula_name(col: str) -> str:
    return re.sub(r"[^A-Za-z0-9_]", "_", col)


def ancova_models(df: pd.DataFrame, group_col: Optional[str], baseline_col: Optional[str], endpoint_col: Optional[str]) -> pd.DataFrame:
    if smf is None or not group_col or not baseline_col or not endpoint_col:
        return pd.DataFrame()
    if group_col not in df.columns or baseline_col not in df.columns or endpoint_col not in df.columns:
        return pd.DataFrame()
    d = df[[group_col, baseline_col, endpoint_col]].dropna().copy()
    if len(d) < 6:
        return pd.DataFrame([{"Model": "ANCOVA", "Reason": "n<6"}])
    d = d.rename(columns={group_col: "Group", baseline_col: "Baseline", endpoint_col: "Endpoint"})
    try:
        m = smf.ols("Endpoint ~ C(Group) + Baseline", data=d).fit(cov_type="HC3")
        rows = []
        for name, coef in m.params.items():
            rows.append({"Model": "ANCOVA", "Term": name, "Estimate": coef, "SE": m.bse[name], "p-value": m.pvalues[name], "R2": m.rsquared, "n": int(m.nobs), "Reason": "OK"})
        return pd.DataFrame(rows)
    except Exception as e:
        return pd.DataFrame([{"Model": "ANCOVA", "Reason": f"Failed: {e}"}])


def factorial_anova(df: pd.DataFrame, outcomes: List[str], factor_cols: Sequence[str]) -> pd.DataFrame:
    if smf is None or anova_lm is None or not factor_cols or len(factor_cols) < 2:
        return pd.DataFrame()
    factor_cols = [c for c in factor_cols if c in df.columns]
    if len(factor_cols) < 2:
        return pd.DataFrame()
    rows = []
    for var in outcomes:
        use_cols = factor_cols + [var]
        d = df[use_cols].dropna().copy()
        if len(d) < 8:
            rows.append({"Variable": var, "Reason": "n<8 for factorial ANOVA"})
            continue
        rename = {c: safe_formula_name(c) for c in use_cols}
        d = d.rename(columns=rename)
        f1, f2 = rename[factor_cols[0]], rename[factor_cols[1]]
        y = rename[var]
        formula = f"{y} ~ C({f1}) * C({f2})"
        try:
            model = smf.ols(formula, data=d).fit()
            aov = anova_lm(model, typ=2)
            for term, r in aov.reset_index().rename(columns={"index": "Term"}).iterrows():
                rows.append({"Variable": var, "Term": r["Term"], "F": r.get("F", np.nan), "p-value": r.get("PR(>F)", np.nan), "Reason": "OK"})
        except Exception as e:
            rows.append({"Variable": var, "Reason": f"Failed: {e}"})
    return pd.DataFrame(rows)


def glm_models(df: pd.DataFrame, outcomes: List[str], group_col: Optional[str]) -> pd.DataFrame:
    if sm is None or smf is None or not group_col or group_col not in df.columns:
        return pd.DataFrame()
    rows = []
    for var in outcomes:
        x = df[var].dropna()
        unique = sorted(pd.unique(x))
        if len(unique) == 2 and set(unique).issubset({0, 1}):
            family = sm.families.Binomial()
            model_type = "Logistic GLM"
        elif (x >= 0).all() and np.all(np.isclose(x, np.round(x))) and x.mean() > 0:
            ratio = x.var(ddof=1) / max(x.mean(), 1e-8)
            family = sm.families.NegativeBinomial() if ratio > 1.5 else sm.families.Poisson()
            model_type = "Negative binomial GLM" if ratio > 1.5 else "Poisson GLM"
        else:
            continue
        d = df[[group_col, var]].dropna().copy().rename(columns={group_col: "Group", var: "Outcome"})
        if len(d) < 8:
            rows.append({"Variable": var, "Model": model_type, "Reason": "n<8"})
            continue
        try:
            model = smf.glm("Outcome ~ C(Group)", data=d, family=family).fit()
            for name, coef in model.params.items():
                rows.append({"Variable": var, "Model": model_type, "Term": name, "Estimate": coef, "SE": model.bse[name], "p-value": model.pvalues[name], "n": int(model.nobs), "Reason": "OK"})
        except Exception as e:
            rows.append({"Variable": var, "Model": model_type, "Reason": f"Failed: {e}"})
    return pd.DataFrame(rows)


def mixed_models(df: pd.DataFrame, outcomes: List[str], group_col: Optional[str], subject_col: Optional[str], period_col: Optional[str], sequence_col: Optional[str], carryover_col: Optional[str], run: bool) -> pd.DataFrame:
    if not run or smf is None or not group_col or not subject_col:
        return pd.DataFrame()
    if group_col not in df.columns or subject_col not in df.columns:
        return pd.DataFrame()
    rows = []
    for var in outcomes:
        cols = [group_col, subject_col, var] + [c for c in [period_col, sequence_col, carryover_col] if c and c in df.columns]
        d = df[cols].dropna().copy()
        if d[subject_col].nunique() < 3 or len(d) < 8:
            rows.append({"Variable": var, "Model": "MixedLM", "Converged": False, "Reason": "Insufficient subjects or observations"})
            continue
        d = d.rename(columns={group_col: "Group", subject_col: "Subject", var: "Outcome"})
        terms = ["C(Group)"]
        if period_col and period_col in cols:
            d = d.rename(columns={period_col: "Period"})
            terms.append("C(Period)")
        if sequence_col and sequence_col in cols:
            d = d.rename(columns={sequence_col: "Sequence"})
            terms.append("C(Sequence)")
        if carryover_col and carryover_col in cols:
            d = d.rename(columns={carryover_col: "Carryover"})
            terms.append("Carryover")
        formula = "Outcome ~ " + " + ".join(terms)
        try:
            with warnings.catch_warnings(record=True) as wlist:
                warnings.simplefilter("always")
                model = smf.mixedlm(formula, data=d, groups=d["Subject"])
                res = model.fit(reml=False, method="lbfgs", maxiter=200, disp=False)
            warn_text = "; ".join(str(w.message) for w in wlist)
            for name, coef in res.params.items():
                if name == "Group Var":
                    continue
                rows.append({"Variable": var, "Model": "MixedLM", "Term": name, "Estimate": coef, "SE": res.bse.get(name, np.nan), "p-value": res.pvalues.get(name, np.nan), "Converged": bool(res.converged), "Warnings": warn_text, "Reason": "OK" if res.converged else "Not converged"})
        except Exception as e:
            rows.append({"Variable": var, "Model": "MixedLM", "Converged": False, "Reason": f"Failed: {e}"})
    return pd.DataFrame(rows)


def correlation_screen(df: pd.DataFrame, outcomes: List[str], max_pairs: int = 2000) -> Tuple[pd.DataFrame, pd.DataFrame]:
    rows_p, rows_s = [], []
    cols = outcomes[:]
    count = 0
    for i in range(len(cols)):
        for j in range(i + 1, len(cols)):
            if count >= max_pairs:
                return pd.DataFrame(rows_p), pd.DataFrame(rows_s)
            a, b = cols[i], cols[j]
            d = df[[a, b]].dropna()
            if len(d) < 3 or d[a].nunique() < 2 or d[b].nunique() < 2:
                continue
            rp, pp = stats.pearsonr(d[a], d[b])
            rs, ps = stats.spearmanr(d[a], d[b])
            rows_p.append({"Variable 1": a, "Variable 2": b, "r": rp, "p-value": pp, "n": len(d)})
            rows_s.append({"Variable 1": a, "Variable 2": b, "rho": rs, "p-value": ps, "n": len(d)})
            count += 1
    return pd.DataFrame(rows_p), pd.DataFrame(rows_s)


def ols_screen(df: pd.DataFrame, outcomes: List[str], group_col: Optional[str]) -> pd.DataFrame:
    if smf is None or not group_col or group_col not in df.columns:
        return pd.DataFrame()
    rows = []
    for var in outcomes:
        d = df[[group_col, var]].dropna().copy().rename(columns={group_col: "Group", var: "Outcome"})
        if len(d) < 6:
            continue
        try:
            m = smf.ols("Outcome ~ C(Group)", data=d).fit(cov_type="HC3")
            for term in m.params.index:
                rows.append({"Variable": var, "Term": term, "Estimate": m.params[term], "SE_HC3": m.bse[term], "p-value": m.pvalues[term], "R2": m.rsquared, "n": int(m.nobs)})
        except Exception as e:
            rows.append({"Variable": var, "Reason": f"Failed: {e}"})
    return pd.DataFrame(rows)


def data_audit(df: pd.DataFrame) -> pd.DataFrame:
    rows = []
    for col in df.columns:
        rows.append({
            "Column": col,
            "dtype": str(df[col].dtype),
            "n": len(df),
            "non_missing": int(df[col].notna().sum()),
            "missing": int(df[col].isna().sum()),
            "unique": int(df[col].nunique(dropna=True)),
            "numeric": bool(pd.api.types.is_numeric_dtype(df[col])),
            "zero_variance": bool(pd.api.types.is_numeric_dtype(df[col]) and df[col].nunique(dropna=True) <= 1),
        })
    return pd.DataFrame(rows)


def set_cell_border(cell, top=None, bottom=None):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = tcPr.first_child_found_in("w:tcBorders")
    if tcBorders is None:
        tcBorders = OxmlElement("w:tcBorders")
        tcPr.append(tcBorders)
    for edge, val in [("top", top), ("bottom", bottom)]:
        if val:
            tag = "w:" + edge
            element = tcBorders.find(qn(tag))
            if element is None:
                element = OxmlElement(tag)
                tcBorders.append(element)
            element.set(qn("w:val"), "single")
            element.set(qn("w:sz"), "8")
            element.set(qn("w:space"), "0")
            element.set(qn("w:color"), "000000")


def style_doc(doc: Document):
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    style.font.size = Pt(11)
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)


def add_three_line_table(doc: Document, table_df: pd.DataFrame, title: str):
    p = doc.add_paragraph()
    run = p.add_run(title)
    run.bold = True
    cols = [c for c in table_df.columns if not c.endswith(" letter")]
    table = doc.add_table(rows=1, cols=len(cols))
    table.autofit = True
    hdr = table.rows[0].cells
    for i, col in enumerate(cols):
        hdr[i].text = str(col)
        for r in hdr[i].paragraphs[0].runs:
            r.bold = True
        set_cell_border(hdr[i], top=True, bottom=True)
    for _, row in table_df.iterrows():
        cells = table.add_row().cells
        for i, col in enumerate(cols):
            val = row.get(col, "")
            if col == "p-value" and not isinstance(val, str):
                val = format_p(val)
            cells[i].text = "" if pd.isna(val) else str(val)
    for cell in table.rows[-1].cells:
        set_cell_border(cell, bottom=True)
    for row in table.rows:
        for cell in row.cells:
            for par in cell.paragraphs:
                par.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in par.runs:
                    run.font.name = "Times New Roman"
                    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
                    run.font.size = Pt(9)
    foot = doc.add_paragraph()
    foot.style = doc.styles["Normal"]
    foot.add_run("Note. ").italic = True
    foot.add_run("Continuous variables are presented as mean ± SEM unless indicated as median (IQR). For tables with three or more groups, different superscript letters within a row indicate significant pairwise differences at p < 0.05. The detailed pairwise table is not displayed by default.")
    doc.add_paragraph("")


def write_word_tables(path: Path, table_df: pd.DataFrame):
    if Document is None:
        return
    doc = Document()
    style_doc(doc)
    doc.add_heading("Three-line Tables", level=1)
    if table_df.empty:
        doc.add_paragraph("No group-comparison table was generated. Check input data and group column.")
    else:
        # Keep readable; split by chunks of variables if needed.
        add_three_line_table(doc, table_df, "Table 1. Group comparison summary")
    doc.save(path)


def write_methods_results(path: Path, df: pd.DataFrame, outcomes: List[str], comp: pd.DataFrame, paired: pd.DataFrame, ancova: pd.DataFrame, factorial: pd.DataFrame, glm: pd.DataFrame, mixed: pd.DataFrame):
    if Document is None:
        return
    doc = Document()
    style_doc(doc)
    doc.add_heading("Statistical Methods and Results Draft", level=1)
    doc.add_heading("Statistical methods", level=2)
    doc.add_paragraph(
        "Data were inspected for missing values, zero variance, and invalid numerical entries before analysis. "
        "Continuous variables were summarized as mean ± SEM, while ordinal or score-like variables were summarized as median (IQR). "
        "Independent two-group comparisons were performed using Welch's t-test or the Mann–Whitney U test as appropriate. "
        "For three or more groups, one-way ANOVA or Kruskal–Wallis tests were used. "
        "Two-sided p-values were reported, with p < 0.05 considered statistically significant. "
        "q-values and FDR q-values were not reported by default. "
        "For multi-group tables, compact letter displays were generated from background pairwise comparisons for table annotation."
    )
    doc.add_heading("Data audit summary", level=2)
    doc.add_paragraph(f"The dataset contained {df.shape[0]} rows and {df.shape[1]} columns. {len(outcomes)} numeric outcome variable(s) were identified for screening.")
    doc.add_heading("Results summary", level=2)
    if not comp.empty and "p-value" in comp.columns:
        sig = comp[(comp["p-value"].notna()) & (comp["p-value"] < 0.05)]
        doc.add_paragraph(f"Group-comparison screening identified {len(sig)} variable(s) with p < 0.05 out of {comp['Variable'].nunique()} tested variable(s).")
        for _, r in sig.head(10).iterrows():
            doc.add_paragraph(f"{r['Variable']}: {r['Test']}, p = {format_p(r['p-value'])}.", style=None)
    else:
        doc.add_paragraph("No valid group comparison was generated.")
    if not paired.empty:
        doc.add_paragraph(f"Paired analysis generated {len(paired)} result row(s).")
    if not ancova.empty:
        doc.add_paragraph(f"ANCOVA module generated {len(ancova)} result row(s).")
    if not factorial.empty:
        doc.add_paragraph(f"Factorial ANOVA module generated {len(factorial)} result row(s).")
    if not glm.empty:
        doc.add_paragraph(f"GLM module generated {len(glm)} result row(s).")
    if not mixed.empty:
        failed = mixed[(mixed.get("Converged", True) == False)] if "Converged" in mixed.columns else pd.DataFrame()
        doc.add_paragraph(f"Mixed model module generated {len(mixed)} result row(s); {len(failed)} row(s) indicated non-convergence or insufficient data.")
    doc.add_heading("Interpretation caution", level=2)
    doc.add_paragraph(
        "These outputs should be interpreted in the context of the experimental design, sample size, measurement units, and biological or substantive plausibility. "
        "Associations from observational or exploratory analyses should not be interpreted as causal effects without appropriate design and modelling support."
    )
    doc.save(path)


def make_figures(df: pd.DataFrame, outcomes: List[str], group_col: Optional[str], outdir: Path, max_figures: int = 12):
    if plt is None or not group_col or group_col not in df.columns:
        return
    figdir = outdir / "figures"
    figdir.mkdir(parents=True, exist_ok=True)
    for var in outcomes[:max_figures]:
        d = df[[group_col, var]].dropna()
        if d.empty:
            continue
        groups = [str(g) for g in pd.unique(d[group_col])]
        data = [d.loc[d[group_col].astype(str) == g, var].astype(float).values for g in groups]
        fig = plt.figure(figsize=(6, 4))
        ax = fig.add_subplot(111)
        ax.boxplot(data, tick_labels=groups, showmeans=True)
        for i, arr in enumerate(data, start=1):
            jitter = np.random.default_rng(1).normal(i, 0.03, len(arr))
            ax.plot(jitter, arr, "o", markersize=3)
        ax.set_title(var)
        ax.set_xlabel(group_col)
        ax.set_ylabel(var)
        fig.tight_layout()
        safe = re.sub(r"[^A-Za-z0-9_.-]+", "_", var)[:80]
        fig.savefig(figdir / f"{safe}.png", dpi=300)
        fig.savefig(figdir / f"{safe}.pdf")
        plt.close(fig)


def write_excel(path: Path, sheets: Dict[str, pd.DataFrame]):
    # Generated by the Skill script. Keep raw numeric precision in cells; display formatting is left to spreadsheet software.
    with pd.ExcelWriter(path, engine="openpyxl") as writer:
        for name, df in sheets.items():
            if df is None or df.empty:
                df = pd.DataFrame({"Message": ["No results generated for this module."]})
            sheet = re.sub(r"[\\/*?:\[\]]", " ", name)[:31]
            df.to_excel(writer, sheet_name=sheet, index=False)
        wb = writer.book
        for ws in wb.worksheets:
            ws.freeze_panes = "A2"
            for cell in ws[1]:
                font = copy(cell.font); font.bold = True; cell.font = font
            for col_cells in ws.columns:
                max_len = max(len(str(c.value)) if c.value is not None else 0 for c in col_cells[:100])
                ws.column_dimensions[col_cells[0].column_letter].width = min(max(10, max_len + 2), 38)


def parse_args(argv=None):
    p = argparse.ArgumentParser(description="Universal research statistics analysis")
    p.add_argument("--input", required=True, help="Input CSV/XLSX file")
    p.add_argument("--outdir", required=True, help="Output directory")
    p.add_argument("--group-col", default=None)
    p.add_argument("--subject-col", default=None)
    p.add_argument("--time-col", default=None)
    p.add_argument("--period-col", default=None)
    p.add_argument("--sequence-col", default=None)
    p.add_argument("--carryover-col", default=None)
    p.add_argument("--baseline-col", default=None)
    p.add_argument("--endpoint-col", default=None)
    p.add_argument("--factor-cols", nargs="*", default=[])
    p.add_argument("--include-posthoc", action="store_true")
    p.add_argument("--run-glm", action="store_true")
    p.add_argument("--run-mixed-models", action="store_true")
    p.add_argument("--make-figures", action="store_true")
    return p.parse_args(argv)


def main(argv=None):
    args = parse_args(argv)
    infile = Path(args.input)
    outdir = Path(args.outdir)
    outdir.mkdir(parents=True, exist_ok=True)
    (outdir / "logs").mkdir(exist_ok=True)

    raw = load_data(infile)
    # Resolve protected columns before coercion so labels are not damaged.
    col_roles = {k: resolve_col(raw, getattr(args, k)) for k in [
        "group_col", "subject_col", "time_col", "period_col", "sequence_col", "carryover_col", "baseline_col", "endpoint_col"
    ]}
    factor_cols = [resolve_col(raw, c) for c in args.factor_cols]
    factor_cols = [c for c in factor_cols if c]
    protected = [c for c in list(col_roles.values()) + factor_cols if c]
    df, coercion_log = numeric_coercion(raw, protected_cols=protected)
    # Re-resolve after coercion.
    group_col = col_roles["group_col"]
    subject_col = col_roles["subject_col"]
    time_col = col_roles["time_col"]
    period_col = col_roles["period_col"]
    sequence_col = col_roles["sequence_col"]
    carryover_col = col_roles["carryover_col"]
    baseline_col = col_roles["baseline_col"]
    endpoint_col = col_roles["endpoint_col"]

    exclude = [group_col, subject_col, time_col, period_col, sequence_col, carryover_col, baseline_col, endpoint_col] + factor_cols
    outcomes = numeric_outcomes(df, exclude=exclude)
    audit = data_audit(df)
    desc = descriptive_by_group(df, outcomes, group_col)
    comp, posthoc, table_data = group_comparisons(df, outcomes, group_col)
    paired = paired_comparisons(df, outcomes, subject_col, time_col)
    ancova = ancova_models(df, group_col, baseline_col, endpoint_col)
    factorial = factorial_anova(df, outcomes, factor_cols)
    glm = glm_models(df, outcomes, group_col) if args.run_glm else pd.DataFrame()
    mixed = mixed_models(df, outcomes, group_col, subject_col, period_col, sequence_col, carryover_col, args.run_mixed_models)
    pearson, spearman = correlation_screen(df, outcomes)
    ols = ols_screen(df, outcomes, group_col)

    run_info = pd.DataFrame([{
        "input": str(infile), "rows": df.shape[0], "columns": df.shape[1],
        "group_col": group_col, "subject_col": subject_col, "time_col": time_col,
        "period_col": period_col, "sequence_col": sequence_col,
        "baseline_col": baseline_col, "endpoint_col": endpoint_col,
        "factor_cols": ", ".join(factor_cols), "outcomes_detected": len(outcomes),
        "p_value_policy": "p-value only; no q-value/FDR q-value by default",
        "compact_letters": "Generated for >=3 groups; detailed pairwise table hidden unless --include-posthoc",
    }])
    column_audit = pd.DataFrame([{"Original column": c, "Resolved role": next((k for k, v in col_roles.items() if v == c), "") or ("factor" if c in factor_cols else "")} for c in raw.columns])
    skipped = []
    if not outcomes:
        skipped.append({"Module": "Outcome detection", "Reason": "No numeric outcome variables detected"})
    if not group_col:
        skipped.append({"Module": "Group comparisons", "Reason": "No group column supplied"})
    if not args.include_posthoc:
        skipped.append({"Module": "Post hoc tests sheet", "Reason": "Detailed post hoc table hidden by default; compact letters retained in Word table"})
    skipped_df = pd.DataFrame(skipped)

    sheets = {
        "Run info": run_info,
        "Column audit": column_audit,
        "Numeric coercion log": coercion_log,
        "Data audit": audit,
        "Descriptive statistics": desc,
        "Group comparisons": comp,
        "Three-line table data": table_data,
        "Paired comparisons": paired,
        "ANCOVA": ancova,
        "Factorial ANOVA": factorial,
        "GLM": glm,
        "Mixed models": mixed,
        "Pearson correlations": pearson,
        "Spearman correlations": spearman,
        "OLS screening": ols,
        "Skipped analyses": skipped_df,
    }
    if args.include_posthoc:
        sheets["Post hoc tests"] = posthoc
    write_excel(outdir / "analysis_results.xlsx", sheets)
    write_word_tables(outdir / "three_line_tables.docx", table_data)
    write_methods_results(outdir / "methods_results.docx", df, outcomes, comp, paired, ancova, factorial, glm, mixed)
    if args.make_figures:
        make_figures(df, outcomes, group_col, outdir)
    coercion_log.to_csv(outdir / "logs" / "numeric_coercion_log.csv", index=False)
    audit.to_csv(outdir / "logs" / "data_audit.csv", index=False)
    table_data.to_csv(outdir / "logs" / "three_line_table_data.csv", index=False)
    print(json.dumps({
        "status": "ok",
        "outputs": ["analysis_results.xlsx", "three_line_tables.docx", "methods_results.docx"],
        "outcomes_detected": len(outcomes),
        "rows": int(df.shape[0]),
        "columns": int(df.shape[1])
    }, ensure_ascii=False, indent=2))


if __name__ == "__main__":
    main()
